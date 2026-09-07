"""PDF + DOCX → plain text. Pure-function; raises classified errors.

PDF strategy: pypdf first (fast, handles most modern PDFs). If the result
looks empty/garbled (heuristic: total length < 50 chars after stripping),
fall back to pdfminer.six (slower, more layout-tolerant). If both fail,
raise ParserError("no_text_extracted").

DOCX strategy: python-docx. Walk paragraphs and table cells, join with newlines.

Legacy .doc (binary Word) is explicitly rejected as
ParserError("doc_legacy_not_supported") — parsing it needs antiword or
LibreOffice (binary deps); deferred to a later plan.
"""

from __future__ import annotations

import io
from typing import Final

import anyio.to_thread
import pypdf
import pypdf.errors
from docx import Document
from pdfminer.high_level import extract_text as pdfminer_extract
from pdfminer.pdfdocument import PDFEncryptionError
from pdfminer.pdfparser import PDFSyntaxError

from jobify.integrations.parser.base import ParserError, TransientParserError

PDF_CONTENT_TYPE: Final[str] = "application/pdf"
DOCX_CONTENT_TYPE: Final[str] = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
LEGACY_DOC_CONTENT_TYPE: Final[str] = "application/msword"

MAX_TEXT_BYTES: Final[int] = 64 * 1024  # 64 KB cap on extracted text
_GARBLED_THRESHOLD: Final[int] = 50  # pypdf result shorter than this → try pdfminer
# Letter-spaced extraction ("G a n e s h") -- a font-encoding artefact where
# pypdf emits every glyph as its own token. Seen on 1 of 12 real resumes
# (2026-09-07): 1,651 single-character tokens cleared the length threshold
# above, so pdfminer (which read the same file cleanly, 7 % single-char
# tokens) never ran. Normal prose sits well under 0.2 (initials, bullets).
_LETTER_SPACED_MIN_TOKENS: Final[int] = 20
_LETTER_SPACED_RATIO: Final[float] = 0.6


async def extract_text(*, content: bytes, content_type: str) -> str:
    """Extract plain text from a resume blob. Truncated to MAX_TEXT_BYTES."""
    if content_type == LEGACY_DOC_CONTENT_TYPE:
        raise ParserError("doc_legacy_not_supported")
    if content_type == PDF_CONTENT_TYPE:
        text = await anyio.to_thread.run_sync(_extract_pdf, content)
        return _truncate(text)
    if content_type == DOCX_CONTENT_TYPE:
        text = await anyio.to_thread.run_sync(_extract_docx, content)
        return _truncate(text)
    raise ParserError("unsupported_content_type")


def _extract_pdf(content: bytes) -> str:
    """Try pypdf; fall back to pdfminer if the result looks empty or garbled.

    Two garbled signatures, checked in order:
    - Empty/near-empty (shorter than _GARBLED_THRESHOLD after stripping):
      image-only PDFs where extractors return near-nothing.
    - Letter-spaced (most tokens are single characters): a font-encoding
      artefact that clears the length check with unusable text. pdfminer
      usually reads these cleanly; whichever extractor produces the lower
      single-character-token ratio wins.
    Non-empty short results (e.g., single-line PDFs) are returned as-is from
    pypdf without triggering the pdfminer fallback.
    """
    pypdf_text = _extract_pdf_pypdf(content)
    pypdf_stripped = pypdf_text.strip()
    if len(pypdf_stripped) >= _GARBLED_THRESHOLD:
        if not _looks_letter_spaced(pypdf_text):
            # pypdf produced substantial, readable text — use it directly.
            return pypdf_text
        pdfminer_text = _extract_pdf_pdfminer(content)
        if pdfminer_text.strip() and _single_char_token_ratio(
            pdfminer_text
        ) < _single_char_token_ratio(pypdf_text):
            return pdfminer_text
        return pypdf_text
    if len(pypdf_stripped) > 0:
        # pypdf produced a short but non-empty result.  For a real resume this
        # would be unusual, but for test fixtures (or trivially short docs) it
        # is valid.  Attempt pdfminer for a potentially better extraction; if
        # pdfminer also returns short text, prefer whichever is longer.
        pdfminer_text = _extract_pdf_pdfminer(content)
        if len(pdfminer_text.strip()) > len(pypdf_stripped):
            return pdfminer_text
        return pypdf_text
    # pypdf returned nothing — try pdfminer as the primary fallback.
    pdfminer_text = _extract_pdf_pdfminer(content)
    if pdfminer_text.strip():
        return pdfminer_text
    # Both extractors returned nothing. Image-only / scanned PDF.
    raise ParserError("no_text_extracted")


def _single_char_token_ratio(text: str) -> float:
    """Share of whitespace-separated tokens that are exactly one character.

    Returns 0.0 for texts too short to judge (fewer than
    _LETTER_SPACED_MIN_TOKENS tokens) so tiny fixtures never trip the check.
    """
    tokens = text.split()
    if len(tokens) < _LETTER_SPACED_MIN_TOKENS:
        return 0.0
    return sum(1 for token in tokens if len(token) == 1) / len(tokens)


def _looks_letter_spaced(text: str) -> bool:
    return _single_char_token_ratio(text) >= _LETTER_SPACED_RATIO


def _extract_pdf_pypdf(content: bytes) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            raise ParserError("password_protected")
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ParserError:
        raise
    except pypdf.errors.PdfReadError as exc:
        # Often "EOF marker not found", malformed xref, etc. — permanent.
        raise ParserError("pdf_read_error") from exc
    except Exception as exc:  # unknown library bug; treat as transient
        raise TransientParserError(f"pypdf_unexpected: {type(exc).__name__}") from exc


def _extract_pdf_pdfminer(content: bytes) -> str:
    try:
        return pdfminer_extract(io.BytesIO(content)) or ""
    except PDFEncryptionError as exc:
        raise ParserError("password_protected") from exc
    except PDFSyntaxError as exc:
        raise ParserError("pdf_syntax_error") from exc
    except Exception as exc:  # unknown library bug; treat as transient
        raise TransientParserError(f"pdfminer_unexpected: {type(exc).__name__}") from exc


def _extract_docx(content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:  # python-docx raises generic Exception variants
        raise ParserError("docx_read_error") from exc

    try:
        lines: list[str] = []
        for para in doc.paragraphs:
            if para.text:
                lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        lines.append(cell.text)
    except ParserError:
        raise
    except Exception as exc:
        raise TransientParserError(f"docx_walk_unexpected: {type(exc).__name__}") from exc

    text = "\n".join(lines)
    if not text.strip():
        raise ParserError("no_text_extracted")
    return text


def _truncate(text: str) -> str:
    """Truncate to MAX_TEXT_BYTES of UTF-8 — never split mid-codepoint."""
    encoded = text.encode("utf-8")
    if len(encoded) <= MAX_TEXT_BYTES:
        return text
    # Decode with errors='ignore' drops any partial codepoint at the cut.
    return encoded[:MAX_TEXT_BYTES].decode("utf-8", errors="ignore")
