"""PDF + DOCX → plain text. Pure-function; raises classified errors.

PDF strategy: pypdf first (fast, handles most modern PDFs). If the result
looks empty/garbled (heuristic: total length < 50 chars after stripping),
fall back to pdfminer.six (slower, more layout-tolerant). If both fail,
raise ParserError("no_text_extracted").

DOCX strategy: python-docx. Walk paragraphs and table cells, join with newlines.

Legacy .doc (binary Word) is explicitly rejected as
ParserError("doc_legacy_not_supported") — parsing it needs antiword or
LibreOffice (binary deps); deferred to a later plan.
"""

from __future__ import annotations

import io
from typing import Final

import anyio.to_thread
import pypdf
import pypdf.errors
from docx import Document
from pdfminer.high_level import extract_text as pdfminer_extract
from pdfminer.pdfdocument import PDFEncryptionError
from pdfminer.pdfparser import PDFSyntaxError

from kpa.integrations.parser.base import ParserError, TransientParserError

PDF_CONTENT_TYPE: Final[str] = "application/pdf"
DOCX_CONTENT_TYPE: Final[str] = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
LEGACY_DOC_CONTENT_TYPE: Final[str] = "application/msword"

MAX_TEXT_BYTES: Final[int] = 64 * 1024  # 64 KB cap on extracted text
_EMPTY_THRESHOLD: Final[int] = 50  # pypdf result shorter than this → try pdfminer


async def extract_text(*, content: bytes, content_type: str) -> str:
    """Extract plain text from a resume blob. Truncated to MAX_TEXT_BYTES."""
    if content_type == LEGACY_DOC_CONTENT_TYPE:
        raise ParserError("doc_legacy_not_supported")
    if content_type == PDF_CONTENT_TYPE:
        text = await anyio.to_thread.run_sync(_extract_pdf, content)
        return _truncate(text)
    if content_type == DOCX_CONTENT_TYPE:
        text = await anyio.to_thread.run_sync(_extract_docx, content)
        return _truncate(text)
    raise ParserError("unsupported_content_type")


def _extract_pdf(content: bytes) -> str:
    """Try pypdf; fall back to pdfminer if the result is empty/garbled."""
    pypdf_text = _extract_pdf_pypdf(content)
    if len(pypdf_text.strip()) >= _EMPTY_THRESHOLD:
        return pypdf_text

    pdfminer_text = _extract_pdf_pdfminer(content)
    if len(pdfminer_text.strip()) >= _EMPTY_THRESHOLD:
        return pdfminer_text

    # Both extractors returned ~nothing. Image-only / scanned PDF.
    raise ParserError("no_text_extracted")


def _extract_pdf_pypdf(content: bytes) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            raise ParserError("password_protected")
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ParserError:
        raise
    except pypdf.errors.PdfReadError as exc:
        # Often "EOF marker not found", malformed xref, etc. — permanent.
        raise ParserError("pdf_read_error") from exc
    except Exception as exc:  # unknown library bug; treat as transient
        raise TransientParserError(f"pypdf_unexpected: {type(exc).__name__}") from exc


def _extract_pdf_pdfminer(content: bytes) -> str:
    try:
        return pdfminer_extract(io.BytesIO(content)) or ""
    except PDFEncryptionError as exc:
        raise ParserError("password_protected") from exc
    except PDFSyntaxError as exc:
        raise ParserError("pdf_syntax_error") from exc
    except Exception as exc:  # unknown library bug; treat as transient
        raise TransientParserError(f"pdfminer_unexpected: {type(exc).__name__}") from exc


def _extract_docx(content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:  # python-docx raises generic Exception variants
        raise ParserError("docx_read_error") from exc

    try:
        lines: list[str] = []
        for para in doc.paragraphs:
            if para.text:
                lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        lines.append(cell.text)
    except ParserError:
        raise
    except Exception as exc:
        raise TransientParserError(f"docx_walk_unexpected: {type(exc).__name__}") from exc

    text = "\n".join(lines)
    if not text.strip():
        raise ParserError("no_text_extracted")
    return text


def _truncate(text: str) -> str:
    """Truncate to MAX_TEXT_BYTES of UTF-8 — never split mid-codepoint."""
    encoded = text.encode("utf-8")
    if len(encoded) <= MAX_TEXT_BYTES:
        return text
    # Decode with errors='ignore' drops any partial codepoint at the cut.
    return encoded[:MAX_TEXT_BYTES].decode("utf-8", errors="ignore")
