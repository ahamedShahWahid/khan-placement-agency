"""Scorer invariants added with the report-only fields (2026-09-07).

- Structured lists score by normalised exact match on a key.
- Report-only fields never move ``overall`` (the number the floors were set on).
- The real set is optional: absent folder → empty list, never an error.
"""

from __future__ import annotations

import json
from pathlib import Path

from jobify.eval.parse_f1 import (
    ALL_FIELDS,
    GATED_FIELDS,
    REPORT_ONLY_FIELDS,
    _certification_keys,
    _education_keys,
    _experience_keys,
    eval_examples,
    load_document_examples,
    real_data_dir,
    score_example,
)
from jobify.integrations.parser.base import (
    CertificationEntry,
    EducationEntry,
    ExperienceEntry,
    ParsedResume,
)


def _parsed(**kw: object) -> ParsedResume:
    return ParsedResume(parser_name="t", raw_text="", **kw)  # type: ignore[arg-type]


def test_experience_key_normalises_case_whitespace_and_trailing_punctuation() -> None:
    parsed = [ExperienceEntry(company="  Acme  Corp. ", title="Senior   Engineer,")]
    expected = [{"company": "acme corp", "title": "senior engineer"}]
    assert _experience_keys(parsed) == _experience_keys(expected) == {"acme corp|senior engineer"}


def test_missing_key_part_renders_as_question_mark_not_empty() -> None:
    """A title-less entry must still be a distinct key, and must not collide
    with an entry that has a title."""
    assert _experience_keys([{"company": "Acme"}]) == {"acme|?"}
    assert _experience_keys([{"company": "Acme", "title": "Dev"}]) == {"acme|dev"}
    assert _experience_keys([{}]) == set()  # nothing to key on → no entry


def test_education_and_certification_keys() -> None:
    assert _education_keys([EducationEntry(institution="IIT Bombay", degree="B.Tech")]) == {
        "iit bombay|b.tech"
    }
    assert _certification_keys([CertificationEntry(name="AWS SAA ")]) == {"aws saa"}


def test_structured_fields_score_as_sets_and_do_not_move_overall() -> None:
    parsed = _parsed(
        name="Priya",
        skills=["python"],
        languages=["English"],
        experience=[ExperienceEntry(company="Acme", title="Dev")],
        education=[EducationEntry(institution="IIT", degree="B.Tech")],
        certifications=[CertificationEntry(name="AWS")],
    )
    expected = {
        "name": "Priya",
        "skills": ["python"],
        # Every report-only expectation deliberately WRONG:
        "languages": ["Hindi"],
        "experience": [{"company": "Other", "title": "Dev"}],
        "education": [],
        "certifications": [{"name": "GCP"}],
    }
    result = score_example("x", parsed, expected)  # type: ignore[arg-type]

    assert set(result.per_field) == set(ALL_FIELDS)
    for fname in REPORT_ONLY_FIELDS:
        assert result.per_field[fname].f1() < 1.0, fname
    # Gated fields are all perfect (email/phone: None vs None), so overall
    # must be 1.0 regardless of the report-only wreckage.
    assert result.overall_f1() == 1.0

    report = eval_examples([("x", "", expected)], lambda _t: parsed)  # type: ignore[list-item]
    assert report.overall_f1 == 1.0
    assert report.per_field_f1["experience"] == 0.0
    assert "[report-only]" in report.summary()
    assert "acme|dev" in report.token_diff()


def test_gated_fields_are_exactly_the_original_four() -> None:
    """Adding a gated field silently changes what 0.85 / 0.90 mean."""
    assert GATED_FIELDS == ("name", "email", "phone", "skills")


def test_real_data_dir_absent_is_none_not_error(monkeypatch, tmp_path: Path) -> None:  # type: ignore[no-untyped-def]
    monkeypatch.setenv("JOBIFY_PARSE_EVAL_REAL_DIR", str(tmp_path / "nope"))
    assert real_data_dir() is None


def test_load_document_examples_skips_documents_without_expected_file(tmp_path: Path) -> None:
    (tmp_path / "a.pdf").write_bytes(b"%PDF-1.4 not really")
    (tmp_path / "b.txt").write_text("ignored: not a document type")
    (tmp_path / "README.md").write_text("ignored")
    # No expected.json for a.pdf → skipped before any extraction is attempted.
    assert load_document_examples(tmp_path) == []


def test_load_document_examples_extracts_docx_text(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Asha Verma")
    doc.add_paragraph("asha@example.com")
    doc.save(tmp_path / "asha.docx")
    (tmp_path / "asha.expected.json").write_text(json.dumps({"name": "Asha Verma"}))

    examples = load_document_examples(tmp_path)

    assert len(examples) == 1
    example_id, text, expected = examples[0]
    assert example_id == "asha"
    assert "asha@example.com" in text
    assert expected == {"name": "Asha Verma"}
